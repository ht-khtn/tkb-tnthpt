#!/usr/bin/env python3
"""Extract timetable entries from the Word timetable document into CSV.

Output columns:
- id
- weekday
- period_no
- location
- subject_name

The timetable table uses columns after the first two as locations/rooms.
Each non-empty cell may contain "<subject> - <teacher>"; only the subject is
kept in the output.
"""

from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path

from docx import Document  # type: ignore


BASE_DIR = Path(__file__).resolve().parents[1]
DEFAULT_INPUT = BASE_DIR / "word" / "TKB ON THI TN.xlsx.docx"
DEFAULT_OUTPUT = BASE_DIR / "csv" / "thpt_timetable.csv"


def clean_text(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split()).strip()


def split_subject(value: str) -> str:
    value = clean_text(value)
    if not value:
        return ""
    parts = re.split(r"\s*-\s*", value, maxsplit=1)
    return clean_text(parts[0]) if parts else value


def extract_entries(doc_path: Path) -> list[dict[str, str]]:
    doc = Document(str(doc_path))
    if not doc.tables:
        raise ValueError(f"No tables found in {doc_path}")

    table = doc.tables[0]
    headers = [clean_text(cell.text) for cell in table.rows[0].cells]
    if len(headers) < 3:
        raise ValueError("Unexpected timetable header structure")

    locations = headers[2:]
    entries: list[dict[str, str]] = []

    for row in table.rows[1:]:
        weekday = clean_text(row.cells[0].text)
        period_text = clean_text(row.cells[1].text)
        if not weekday or not period_text:
            continue

        try:
            period_no = int(period_text)
        except ValueError:
            continue

        for location, cell in zip(locations, row.cells[2:]):
            subject_name = split_subject(cell.text)
            if not subject_name:
                continue

            entries.append(
                {
                    "weekday": weekday,
                    "period_no": str(period_no),
                    "location": location,
                    "subject_name": subject_name,
                }
            )

    for idx, entry in enumerate(entries, 1):
        entry["id"] = str(idx)

    return entries


def write_csv(output_path: Path, entries: list[dict[str, str]]) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["id", "weekday", "period_no", "location", "subject_name"],
        )
        writer.writeheader()
        writer.writerows(entries)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract timetable CSV from Word")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    entries = extract_entries(args.input)
    write_csv(args.output, entries)
    print(f"Extracted {len(entries)} rows to {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())